# а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
# б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
# в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер
# шрифта.

# а)

from docx import Document
from docx.shared import Pt

doc = Document()
p = doc.add_paragraph()

run1 = p.add_run("Hello ")
run1.bold = False
run2 = p.add_run("Python")
run2.bold = True

doc.save("input.docx")

# б)

doc = Document("input.docx")

bold_text = []

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text.append(run.text)

result = "".join(bold_text)
print(result)

# в)

new_doc = Document()
p = new_doc.add_paragraph()
run = p.add_run("New paragraph for the task")

run.font.name = "Georgia"
run.font.size = Pt(16)
run.italic =True

new_doc.save("output.docx")