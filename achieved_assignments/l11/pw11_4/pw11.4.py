# а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
# б) Прочитайте этот файл в python-строку и выведите на экран.
# в) Создайте абзац с текстом и запишите это в новый word-файл.

from docx import Document



doc = Document()
doc.add_paragraph("Hello Python")
doc.save("input.docx")

doc = Document("input.docx")



text = []

for paragraph in doc.paragraphs:
    text.append(paragraph.text)

result = "\n".join(text)

print(result)



new_text = (
    "Это абзац текста, "
    "который уже существует внутри программы. "
    "Он может быть длинным, состоять из нескольких предложений "
    "и логически представлять один абзац."
)

new_doc = Document()
new_doc.add_paragraph(new_text)
new_doc.save("new.docx")